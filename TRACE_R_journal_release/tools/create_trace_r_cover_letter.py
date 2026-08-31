"""Create the TRACE-R journal cover letter as a polished Word document.

Author: Amir Ghorbani <amir.ghorbani@rmit.edu.au>
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docx" / "TRACE_R_cover_letter_IEEE_TITS_20260831.docx"

BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 101, 109)
BLACK = RGBColor(0, 0, 0)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, *, size: float = 10, bold: bool = False, color: RGBColor = BLACK) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_body(doc: Document, text: str, *, after: float = 3) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.03
    set_font(paragraph.add_run(text))


def add_metadata(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    set_font(paragraph.add_run(f"{label}: "), bold=True)
    set_font(paragraph.add_run(value))


def build() -> Path:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.60)
    section.bottom_margin = Inches(0.60)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.03

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("TRACE-R | Manuscript submission"), size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Amir Ghorbani | RMIT University"), size=8.5, color=MUTED)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(3)
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("COVER LETTER"), size=14, bold=True, color=BLUE)

    add_metadata(document, "To", "Editor-in-Chief, IEEE Transactions on Intelligent Transportation Systems")
    add_metadata(document, "From", "Amir Ghorbani, School of Engineering, RMIT University")
    add_metadata(document, "Date", "31 August 2026")
    add_metadata(document, "Re", "TRACE-R: Telemetry-Conditioned Restoration for Road-Defect Detection")

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)

    add_body(document, "Dear Editor-in-Chief,")
    add_body(
        document,
        "On behalf of my coauthors, I submit the manuscript \"TRACE-R: Telemetry-Conditioned Restoration for Road-Defect Detection\" for consideration as a regular paper.",
    )
    add_body(
        document,
        "Road-inspection vehicles record camera settings, inertial motion, timing, and vehicle state alongside each image. TRACE-R converts these measurements into cause-specific capture evidence, reconciles that evidence with the observed image, and conditions a multi-degradation restorer according to sensor availability and reliability. The deployed system returns one restored image to an unchanged road-defect detector.",
    )
    add_body(
        document,
        "The evaluation combines matched controlled studies on IVCNZ and PCM with a full-resolution field study using synchronized Sony ILX-LR1 and SBG records from the Collingwood Road Inspection Dataset. The field study uses 320 human-reviewed frames under a temporally disjoint 180/60/80 calibration, validation, and one-time test protocol. Reviewed defect boxes support field calibration where paired clean targets are unavailable. Within each experiment, restorers use matched data and update budgets, selection is confined to validation data, and every method is evaluated with the same frozen detector. TRACE-R provides the strongest deployable mean detection result on both controlled benchmarks and the highest primary field mAP50. Capture-record interventions further verify the role of measured exposure context.",
    )
    add_body(
        document,
        "The manuscript substantially extends the preliminary RMR-P arXiv study through exposure-aligned camera-IMU records, partial-record reliability, image-sensor state reconciliation, hierarchy-wide adapters, matched task-aware training, sealed controlled tests, and a field-calibration procedure for surveys without paired references. The result is a reproducible restoration interface between instrumented vehicles and existing pavement detectors, directly aligned with intelligent transportation perception and infrastructure monitoring.",
    )
    add_body(
        document,
        "Code, configurations, frozen evaluation protocols, and release material are available at https://github.com/AmirNetwork/rmrnet-restoration/tree/trace-r-journal/TRACE_R_journal_release. The CRID image, telemetry, synchronization, and annotation release is being maintained through the same project after privacy review.",
    )
    add_body(
        document,
        "This manuscript is original, is not under consideration elsewhere, and has been approved by all authors. The authors declare no conflict of interest. This work was supported by Australia's National Road Safety Action Grants Program [Grant number: XXXXXX].",
    )
    add_body(document, "Thank you for considering this work.", after=5)
    add_body(document, "Sincerely,", after=2)
    add_body(document, "Amir Ghorbani", after=0)
    add_body(document, "School of Engineering, RMIT University", after=0)
    add_body(document, "amir.ghorbani@rmit.edu.au", after=0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
